import csv
from sklearn.cluster import MiniBatchKMeans
import matplotlib
matplotlib.use('TkAgg')
import matplotlib.pyplot as plt
import pickle
import numpy as np
import sklearn.metrics as metrics
from yellowbrick.cluster import InterclusterDistance
from scipy.optimize import curve_fit
import umap.umap_ as umap
from colorsys import hls_to_rgb
from pylab import *
from datetime import datetime
import os
import argparse
import scipy.stats as scist
from docx import Document
from feature_extraction import LemmaStemmerTokenizer

# Allow for larger CSV files
maxInt = sys.maxsize
while True:
    # decrease the maxInt value by factor 10 
    # as long as the OverflowError occurs.
    try:
        csv.field_size_limit(maxInt)
        break
    except OverflowError:
        maxInt = int(maxInt/10)

def get_clusters(selected_k, data_file, processed_file, centers, years, save_folder="", save=True):
    """

    Parameters
    ----------
    selected_k : selected number of clusters
    data_file : pickle with raw data as list of dictionaries
    processed_file : pickle with transformed data as array
    centers : array. initial centroids from LDA. Can be initialized as 'k-means++'
    years : list of strings. years for intracluster analysis
    save_folder : string. directory to save result, the default is "".
    save : boolean

    Returns
    -------
    output : dictionary. Keys:
        "yr_avg_cost": List of lists. Average funding by year for each cluster.
        "yr_total_cost": List of lists. Total funding by year for each cluster.
        "size": List. Size of each cluster.
        "data_by_cluster": List of lists of dictionaries. Points in each cluster: [ [{Cluster1pt1}, {Cluster1pt2},...], [{Cluster2pt1}, {Cluster2pt2},...], ...]
        "centroids": 10 x K array of cluster centroids,
        "score": List. Silhouette score by cluster
        "model": MiniBatchKMeans model
        "labels": Cluster labels of data points (ordered)

    """
    # Load data as list of dictionaries
    data = pickle.load(open(data_file,"rb"))

    # Transformed data
    X_transformed = pickle.load(open(processed_file,"rb"))

    # Perform mini batch k means
    km = MiniBatchKMeans(n_clusters=selected_k, init=centers, verbose=0, max_no_improvement=None)
    clusters = km.fit_predict(X_transformed)
    scores = metrics.silhouette_samples(X_transformed, clusters)

    # Output data
    cluster_all = []
    costs = []
    yoy = []
    size = []
    mechanisms = []
    for i in range(6): # initialization
        mechanisms.append([])
    MECH_NAMES = "R01", "U01", "R44", "U24", "R21", "U54"

    for i in range(0,selected_k):

        # indices of cluster k
        cluster = [idx for idx, element in enumerate(clusters) if element == i]

        # get points
        cluster_data = [data[ind] for ind in cluster]
        cluster_scores = [scores[ind] for ind in cluster]
        for i in range(len(cluster_data)):
            cluster_data[i]["score"] = cluster_scores[i]
        cluster_all.append(cluster_data)

        # calculate average cost and std
        try:
            average_cost = sum([item["award_amount"] for item in cluster_data])/len(cluster_data)
        except:
            average_cost = 0
        costs.append(average_cost)

        cost_trend = []
        for year in years:
            year_data = [data[ind]["award_amount"] for ind in cluster if data[ind]["year"] == year]
            if len(year_data) == 0:
                cost_trend.append(0)
            else:
                year_cost = sum(year_data) # /len(year_data)
                cost_trend.append(year_cost)

        yoy.append(cost_trend)

        size.append(len(cluster))

        # get number of awards per mechanism
        if len(cluster_data) != 0:
            for j in range(len(mechanisms)):
                mech = len([ind for ind in cluster if data[ind]["mechanism"] == MECH_NAMES[j]])/len(cluster_data)
                mechanisms[j].append(mech)
        else:
            for j in range(len(mechanisms)):
                mechanisms[j].append(0)

    # Get centroids
    # Identify the top terms for each cluster, using the TF-IDF terms with the highest values in the centroid
    order_centroids = km.cluster_centers_.argsort()[:, ::-1]
    vectorizer = pickle.load(open("data/vectorizer.pkl","rb"))
    terms = vectorizer.get_feature_names_out()
    centroids = []
    for i in range(selected_k):
        centroid_list = []
        for ind in order_centroids[i, :15]:
            centroid_list.append(terms[ind])
        centroids.append(centroid_list)

    # Save centroids
    if save:
        centroid_file = open("{}/centroids".format(save_folder), "w", encoding='utf8')
        for i in range(selected_k):
            centroid_file.write("Cluster %d:" % i)
            for ind in order_centroids[i, :15]:
                centroid_file.write(" %s" % terms[ind])
            centroid_file.write("\n")
        centroid_file.close()

    # get scores
    score = metrics.silhouette_score(X_transformed, km.labels_)

    output = {
        "yr_avg_cost": costs, # Average award size by year by cluster
        "yr_total_cost": yoy, # Total award size by year by cluster
        "size": size, # Number of awards in each cluster
        "data_by_cluster": cluster_all,
        "centroids": centroids,
        "score": score, # Silhouette score for
        "model": km, # K-means model
        "labels": clusters, # Ordered list of cluster number labels for each award
        "mechanisms": mechanisms # List of lists: [r01, u01, r44, u24, r21, u54]. Each internal list has number of awards per mechanism by cluster
        }
    return output

def umap_visualization(X_transformed, cluster_labels, silhouette_scores, sizes, save_folder=""):
    #outlier_scores = sklearn.neighbors.LocalOutlierFactor(contamination=0.1).fit_predict(X_transformed)
    #X_transformed = X_transformed[outlier_scores != -1]
    #cluster_labels = cluster_labels[outlier_scores != -1]
    # product = [silhouette_scores[i]*sizes[i] for i in range(len(sizes))]
    
    top_clusters = sorted(range(len(silhouette_scores)), key=lambda i: silhouette_scores[i], reverse=True)[:9]
    n_subset = len(cluster_labels)
    selected_cells = np.random.choice(np.arange(X_transformed.shape[0]), size = n_subset, replace = False)
    mapper = umap.UMAP(metric='hellinger', random_state=42).fit(X_transformed[selected_cells,:])

    embedding = mapper.transform(X_transformed[selected_cells,:])

    # Colors
    colors = ['tab:blue', 'tab:orange', 'tab:green', 'tab:red', 'tab:purple', 'tab:brown', 'tab:pink', 'tab:olive', 'tab:cyan']
    selected_colors = []
    for point in selected_cells:
        if cluster_labels[point] in top_clusters:
            selected_colors.append(colors[top_clusters.index(cluster_labels[point])])
        else:
            selected_colors.append('tab:gray')
    
    # Plot Clusters on UMAP
    plt.figure()
    plt.grid(visible=None)
    plt.scatter(embedding[:, 0], embedding[:, 1], cmap='Spectral', s=5, c=selected_colors)
    plt.gca().set_aspect('equal', 'datalim')
    num_clust = len(np.unique(cluster_labels[selected_cells]))
    #plt.colorbar(boundaries=np.arange(num_clust+1)-0.5).set_ticks(np.arange(num_clust))
    plt.title('UMAP Projection of Awards, TF-IDF', fontsize=14)
    plt.xlabel("UMAP 1")
    plt.ylabel("UMAP 2")

    manager = plt.get_current_fig_manager()
    manager.resize(*manager.window.maxsize())
    plt.savefig('{}/umap.png'.format(save_folder))

def rainbow_color_stops(n=10, end=1, shade=0.9):
    return [ hls_to_rgb(end * i/(n-1)*shade, 0.5*shade, 1*shade) for i in range(n) ]

def get_funding_projections(data):

    # 1. Determine dimensions for plot
    k = len(data["size"])
    factors = []
    for i in range(1, k+1):
        if k / i == i:
            factors.extend([i,i])
        elif k % i == 0:
            factors.append(i)
    dim1, dim2 = factors[int(len(factors)/2)], factors[int(len(factors)/2-1)]

    # 2. Create plot
    fig, axs = plt.subplots(dim1, dim2, sharex='all', sharey='all')

    # 3. Create hidden frame for shared labels
    fig.add_subplot(111, frameon=False)
    plt.grid(visible=None)
    plt.tick_params(labelcolor='none', which='both', top=False, bottom=False, left=False, right=False)
    plt.xlabel("Years from 1985")
    plt.ylabel("Funding ($100 millions)")

    # 4. Get projections
    years_int = list(range(0,36))
    projection = []
    growth = []
    bounds = []
    for i in range(len(data["yr_total_cost"])):
        popt, pcov = curve_fit(lambda t,a,b: a*np.exp(b*t),  years_int,  data["yr_total_cost"][i],  p0=(4000, 0.1))
        std = np.sqrt(np.diagonal(pcov))
        x = np.linspace(0,21,400)
        # upper0 = popt[0]+1.96*std[0]
        # lower0 = popt[0]-1.96*std[0]
        upper1 = popt[1]+1.96*std[1]
        lower1 = popt[1]-1.96*std[1]

        ypred = [popt[0]*np.exp(popt[1]*point) for point in x] #-popt[0]
        projection.append(ypred[-1])
        growth.append(popt[1])
        bounds.append([lower1, upper1])
        
        # projection.append(0)
        # growth.append(0)
        # bounds.append([0,0])

    # 5. Return 2021 projections and growth rate
    return projection, growth, bounds

def viz_centroids(data):
    model = data["model"]
    X_transformed = pickle.load(open("data/processed-data.pkl","rb"))
    plt.figure()
    visualizer = InterclusterDistance(model, random_state=0)
    visualizer.fit(X_transformed)     # Fit the data to the visualizer
    visualizer.show()        # Finalize and render the figure

def predict_clusters(test_data, selected_k, model):
    test_data = pickle.load(open(test_data,"rb"))
    vectorizer = pickle.load(open("data/vectorizer.pkl","rb"))
    input_text = [item["text"] for item in test_data]
    if len(input_text) == 0:
        return [0 for i in range(0,selected_k)], 0
    test_transformed = vectorizer.transform(input_text)
    years = [str(i) for i in range(1985,2021)]
    labels = model.predict(test_transformed)

    # Output data
    cluster_all = []
    costs = []
    yoy = []
    size = []

    for i in range(0,selected_k):

        # indices of cluster k
        cluster = [idx for idx, element in enumerate(labels) if element == i]

        # get points
        cluster_data = [test_data[ind] for ind in cluster]
        cluster_all.append(cluster_data)

        # calculate average cost and std
        try:
            average_cost = sum([item["award_amount"] for item in cluster_data])/len(cluster_data)
        except:
            average_cost = 0
        costs.append(average_cost)

        cost_trend = []
        for year in years:
            year_data = [test_data[ind]["award_amount"] for ind in cluster if test_data[ind]["year"] == year]
            if len(year_data) == 0:
                cost_trend.append(0)
            else:
                year_cost = sum(year_data)
                cost_trend.append(year_cost)

        yoy.append(cost_trend)

        size.append(len(cluster))

    return cluster_all, size

def get_best_cluster(selected_k, num_trials, centers, years, save_folder="", save=True):
    scores = []
    results = {}
    print("Optimizing model...")
    for i in range(num_trials):
        # Generate clusters for a selected k
        data = get_clusters(selected_k, "data/data.pkl", "data/processed-data.pkl", 'k-means++', years, save_folder, save=save)
        j = 0
        for thing in data["data_by_cluster"]:
            for item in thing:
                try:
                    results[item["id"]].append(centroids[j])
                except:
                    results[item["id"]] = [item["id"],item["title"],item["award_amount"],data["centroids"][j]]
            j+=1
        print("Trial {}: Score = {:.3f}".format(str(i+1), data["score"]))
        scores.append(data["score"])
        if data["score"] >= max(scores):
            chosen = data

    return chosen, scores

def get_citations(clusters):
    """

    Parameters
    ----------
    clusters : nested lists of dictionaries representing each award in a cluster.

    Returns
    -------
    total_citations : list of total citations by cluster
    total_papers : list of total papers by cluster
    apts: average APT [0.9, ...]
    lower: lower bound of 95% CI of average APT: "APT (lower - upper)" [0.85,...]
    upper: upper bound of 95% CI of average APT [0.95,...] - "0.9 (0.85-0.95)"

    """

    # Get clusters by project number
    clusters_by_project = []
    for cluster in clusters:
        cluster = [item["project_number"] for item in cluster]
        cluster = list(set(cluster)) # Remove duplicates
        clusters_by_project.append(cluster)

    # Get number of citations, apt, and publication year by paper
    output = {}
    with open("data/citations.csv", newline='', encoding='utf8') as csvfile:
       raw_data = list(csv.reader(csvfile))
       for i in range(1,len(raw_data)): # "rcr": float(raw_data[i][6]),
           output[raw_data[i][0]] = {
               "citations": int(raw_data[i][13]), 
               "apt": float(raw_data[i][11]),
               "year": int(raw_data[i][1])}

    # Get project number and year by paper
    with open("data/publications.csv", newline='', encoding='utf8') as csvfile:
       raw_data = list(csv.reader(csvfile))
       for i in range(1,len(raw_data)):
           if raw_data[i][1] in output.keys():
               output[raw_data[i][1]]["project"] = raw_data[i][0]

    # Calculate total number of citations, total number of papers, average RCR, average APT for each cluster
    total_citations = []
    total_papers = []
    apts = []
    apts_95 = []
    lower = []
    upper = []
    total_availability = []
    # rcrs = []
    for cluster in clusters_by_project:
        cluster_citations = []
        # cluster_rcr = []
        cluster_apt = []
        num_papers = 0
        availability = []

        for idd in cluster:
            papers = [output[key]["citations"] for key in output if output[key]["project"]==idd] # list of all papers associated with cluster by citation count
            # rcr = [output[key]["rcr"] for key in output if output[key]["project"]==idd]
            apt = [output[key]["apt"] for key in output if output[key]["project"]==idd]

            avail_years = [max(0, 2021-output[key]["year"]) for key in output if output[key]["project"]==idd]

            # cluster_rcr.extend(rcr)
            cluster_apt.extend(apt)
            num_papers += len(papers)
            cluster_citations.append(sum(papers))
            availability.append(sum(avail_years))

        total_citations.append(sum(cluster_citations))
        total_papers.append(num_papers)
        apts_95.append(sum([1 for i in cluster_apt if i==0.95])/len(cluster_apt))
        apts.append(np.mean(cluster_apt))

        #create 95% confidence interval for population mean weight
        apts_interval = scist.norm.interval(alpha=0.95, loc=np.mean(cluster_apt), scale=scist.sem(cluster_apt))
        lower.append(apts_interval[0])
        upper.append(apts_interval[1])
        # rcrs.append(sum(cluster_apt)/len(cluster_apt))

        total_availability.append(int(sum(availability)))

    return total_citations, total_papers, apts_95, apts, lower, upper, total_availability

def get_rep_clusters(result):
    path, dirs, files = next(os.walk('{}/clusters'.format(result)))
    file_count = len(files)
    if file_count == 0:
        return
    document = Document()

    for i in range(file_count):
        unique_awards = {}

        # open file
        with open('{}/clusters/cluster-{}.csv'.format(result, str(i)), newline='', encoding='utf8') as csvfile:
            raw_data = list(csv.reader(csvfile))
            for j in range(1,len(raw_data)):
                title = raw_data[j][1]
                organization = raw_data[j][6]
                mechanism = raw_data[j][7]
                year = int(raw_data[j][8])
                score = float(raw_data[j][11])

                # If this is a new title
                if title not in unique_awards:
                    unique_awards[title] = {
                        "organization": organization,
                        "activity": mechanism,
                        "year": year,
                        "score": score,
                        }

                # If the title is already there
                else:
                    current_year = unique_awards[title]["year"]
                    # Use the most recent one
                    if year > current_year:
                        unique_awards[title] = {
                        "organization": organization,
                        "activity": mechanism,
                        "year": year,
                        "score": score,
                        }

        unique_awards_sorted = dict(sorted(unique_awards.items(), key = lambda item: -item[1]["score"]))
        unique_awards_list = list(unique_awards_sorted.items())[0:5]

        p = document.add_paragraph()
        p.add_run('Cluster {}:'.format(str(i))).bold = True
        table = document.add_table(rows=6, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Title'
        hdr_cells[1].text = 'Awardee'
        hdr_cells[2].text = 'Award Activity'
        hdr_cells[3].text = 'Year'
        hdr_cells[4].text = 'Sample Silhouette Score'

        for i in range(len(unique_awards_list)):
            table.cell(i+1,0).text = unique_awards_list[i][0] # Title
            table.cell(i+1,1).text = unique_awards_list[i][1]['organization'] # Awardee
            table.cell(i+1,2).text = unique_awards_list[i][1]['activity'] # Award Activity
            table.cell(i+1,3).text = str(unique_awards_list[i][1]['year']) # Year
            table.cell(i+1,4).text = "{:.2g}".format(unique_awards_list[i][1]['score']) # Sample Silhouette Score

        document.add_page_break()

    document.save('{}/supp_info.docx'.format(result))

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument(
        '--k',
        type=int,
        required=True,
        help='number of clusters',
        default=30,
        )
    parser.add_argument(
        '--trials',
        type=int,
        required=True,
        help='number of trials',
        default=50,
        )
    FLAGS, unparsed = parser.parse_known_args()


    years = [str(i) for i in range(1985,2021)]
    selected_k = FLAGS.k
    num_trials = FLAGS.trials
    centers = 'k-means++'

    # Create folder to save results
    now = datetime.now()
    save_folder = "results/"+now.strftime("%m-%d-%Y--%H%M%S")
    os.mkdir(save_folder)

    # Get best clustering
    data, scores = get_best_cluster(selected_k, num_trials, centers, years, save_folder)
    with open("{}/model_clustering.pkl".format(save_folder), 'wb') as handle:
        pickle.dump(data, handle)

    # Final cluster files
    num = 0
    os.mkdir(save_folder+"/clusters")
    for cluster in data["data_by_cluster"]:
        if cluster == []:
            continue
        keys = cluster[0].keys()
        with open('{}/clusters/cluster-{}.csv'.format(save_folder,str(num)), 'w', newline='', encoding='utf8')  as output_file:
            dict_writer = csv.DictWriter(output_file, keys)
            dict_writer.writeheader()
            dict_writer.writerows(cluster)
        num+=1

    # Silhouette score by cluster
    print("")
    print("------Silhouette scores------")
    X_transformed = pickle.load(open("data/processed-data.pkl","rb"))
    scores = metrics.silhouette_samples(X_transformed, data["labels"])
    tabulated = []
    pairs = [(scores[i],data["labels"][i]) for i in range(len(scores))]
    for i in range(selected_k):
        avg_score = np.mean([j[0] for j in pairs if j[1] == i])
        print("Cluster {}: {}".format(str(i), str(avg_score)))
        tabulated.append(avg_score)
    print("----------------------------")
    print("")

    # Final centroids
    order_centroids = data["model"].cluster_centers_.argsort()[:, ::-1]
    vectorizer = pickle.load(open("data/vectorizer.pkl","rb"))
    terms = vectorizer.get_feature_names_out()
    centroids = []
    centroid_file = open("{}/centroids".format(save_folder), "w", encoding='utf8')
    for i in range(selected_k):
        centroid_file.write("Cluster %d:" % i)
        centroid_list = []
        for ind in order_centroids[i, :15]:
            centroid_file.write(" %s," % terms[ind])
            centroid_list.append(terms[ind])
        centroids.append(centroid_list)
        centroid_file.write("\n")
    centroid_file.close()

    # UMAP Visualization
    X_transformed = pickle.load(open("data/processed-data.pkl","rb"))
    umap_visualization(X_transformed, data["labels"], tabulated, data["size"], save_folder)

    # Get 2021 projections, projected growth rates, and confidence bounds on growth rates by cluster
    projection, growth, bounds = get_funding_projections(data) # 2021 prediction

    # Get 2021 clusters
    model = data["model"]
    clusters_test, size_test = predict_clusters("data/test-data.pkl", selected_k, model)
    x = np.arange(selected_k)
    if size_test == 0:
        cluster_cost_2021 = [0 for i in range(0, selected_k)]
    else:
        cluster_cost_2021 = [(sum([item["award_amount"] for item in group]) if len(group) > 0 else 0) for group in clusters_test]

    # Save 2021 clusters
    num = 0
    os.mkdir("{}/clusters_test".format(save_folder))
    for cluster in clusters_test:
        try:
            keys = cluster[0].keys()
        except:
            num+=1
            continue
        with open('{}/clusters_test/cluster-{}.csv'.format(save_folder,str(num)), 'w', newline='', encoding='utf8')  as output_file:
            dict_writer = csv.DictWriter(output_file, keys)
            dict_writer.writeheader()
            dict_writer.writerows(cluster)
        num+=1

    # Citations and papers
    citations, papers, apt_pct, apt, lower, upper, availability = get_citations(data["data_by_cluster"])

    # Total funding
    total_cluster_funding = [sum([item["award_amount"] for item in group]) for group in data["data_by_cluster"]]

    # Get representative clusters for supp info
    get_rep_clusters(save_folder)

    # All data - note blank columns for description, category
    output = [["Cluster", "Size", "Total", "Citations", "APT % over 95%", "Avg. APT", "95%CI L", "95%CI U", "Papers", "Citations per $1mil funding", "Years of Availability", "Citations per thousand dollars of funding per year", "Projected 2021 Award", "Actual 2021 Award To Date", "Growth Rate", "95%CI L", "95%CI U", "Score", "Description", "Category", "Clinical/Technical", "Centroids", "%R01", "%U01", "%R44", "%U24", "%R21", "%U54"]]
    for i in range(selected_k):
        output.append([i, data["size"][i], total_cluster_funding[i], citations[i], apt_pct[i], apt[i], lower[i], upper[i], papers[i], citations[i]/total_cluster_funding[i]*1e6, availability[i], citations[i]/total_cluster_funding[i]*1e3/availability[i], projection[i], cluster_cost_2021[i], growth[i], bounds[i][0], bounds[i][1], tabulated[i], " ", " ", " ", centroids[i], data["mechanisms"][0][i], data["mechanisms"][1][i], data["mechanisms"][2][i], data["mechanisms"][3][i], data["mechanisms"][4][i], data["mechanisms"][5][i]])
    with open('{}/final_data.csv'.format(save_folder), 'w', newline='', encoding='utf8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerows(output)

    print("Complete.")
